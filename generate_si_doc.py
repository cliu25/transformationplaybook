"""
Generate systems-integration-chapter_V7.docx from the current HTML source.
All hyperlinks preserved. Run: python3 aift-playbook-site-v16/generate_si_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── helpers ───────────────────────────────────────────────────────────────────
def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink'); rPr.append(rStyle)
    if bold: rPr.append(OxmlElement('w:b'))
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)

def color_run(run, hex_color):
    rPr = run._r.get_or_add_rPr()
    c = OxmlElement('w:color'); c.set(qn('w:val'), hex_color); rPr.append(c)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); color_run(r, '161616')

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); color_run(r, '161616')

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); color_run(r, '161616')

def prose(text, small=False, muted=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.size = Pt(10 if small else 11)
    if italic: r.italic = True
    if muted: color_run(r, '525252')
    return p

def label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(8); color_run(r, '6f6f6f')

def callout(label_text, body_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(label_text.upper() + '  '); r1.bold = True; r1.font.size = Pt(8); color_run(r1, '5c2eb8')
    r2 = p.add_run(body_text); r2.font.size = Pt(11)

def info_callout(label_text, items):
    """items: list of strings or tuples (bold_prefix, rest, optional_links)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label_text.upper()); r.bold = True; r.font.size = Pt(8); color_run(r, '6f6f6f')
    for item in items:
        pb = doc.add_paragraph(style='List Bullet')
        pb.paragraph_format.left_indent = Inches(0.5); pb.paragraph_format.space_after = Pt(3)
        if isinstance(item, str):
            pb.add_run(item).font.size = Pt(11)
        elif isinstance(item, dict):
            if item.get('bold'): pb.add_run(item['bold']).bold = True; pb.runs[-1].font.size = Pt(11)
            if item.get('text'): pb.add_run(item['text']).font.size = Pt(11)
            if item.get('links'):
                for url, txt in item['links']:
                    pb.add_run('  '); add_hyperlink(pb, url, txt)

def painpoint(title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run('⚠ PAIN POINT  '); r.bold = True; r.font.size = Pt(8); color_run(r, 'ba7517')
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.25); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(title); r2.bold = True; r2.font.size = Pt(11); color_run(r2, '633806')
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.25); p3.paragraph_format.space_after = Pt(10)
    r3 = p3.add_run(body); r3.font.size = Pt(11); color_run(r3, '5a3e00')

def ready_banner(heading, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('✓  READY: ' + heading); r.bold = True; r.font.size = Pt(11); color_run(r, '1a6b32')
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12); p2.paragraph_format.left_indent = Inches(0.25)
    r2 = p2.add_run(body); r2.font.size = Pt(11); color_run(r2, '525252')

def phase_tag(phase, steps_context):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(4)
    colors = {'Engage': '0043a4', 'Discover': '5c2eb8', 'Execute': '1a6b32'}
    r1 = p.add_run(f'[{phase.upper()}]  '); r1.bold = True; r1.font.size = Pt(9); color_run(r1, colors.get(phase, '525252'))
    r2 = p.add_run('Systems Integration · ' + steps_context); r2.font.size = Pt(9); color_run(r2, '525252')
    doc.add_paragraph('─' * 72).paragraph_format.space_after = Pt(0)

def step_num_label(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'Step {num}  '); r1.bold = True; r1.font.size = Pt(10); color_run(r1, '8a3ffc')
    r2 = p.add_run(text); r2.bold = True; r2.font.size = Pt(11)

def case_study(context, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('CASE STUDY  '); r1.bold = True; r1.font.size = Pt(8)
    r2 = p.add_run(context); r2.bold = True; r2.font.size = Pt(10); color_run(r2, '5c2eb8')
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.25); p2.paragraph_format.space_after = Pt(10)
    r3 = p2.add_run(body); r3.font.size = Pt(11); color_run(r3, '525252')

def section_rule():
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def bullet(parts, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    for part in (parts if isinstance(parts, list) else [(parts, False)]):
        if isinstance(part, str):
            p.add_run(part).font.size = Pt(11)
        elif isinstance(part, tuple) and part[0] == 'bold':
            r = p.add_run(part[1]); r.bold = True; r.font.size = Pt(11)
        elif isinstance(part, tuple) and part[0] == 'link':
            add_hyperlink(p, part[1], part[2])

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(t.rows[0].cells):
        h.text = headers[i]
        h.paragraphs[0].runs[0].bold = True; h.paragraphs[0].runs[0].font.size = Pt(10)
        h.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, cd in enumerate(row):
            p = cells[ci].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            if isinstance(cd, list):
                for frag in cd:
                    if isinstance(frag, tuple) and frag[0] == 'link': add_hyperlink(p, frag[1], frag[2])
                    elif isinstance(frag, tuple) and frag[0] == 'bold': r = p.add_run(frag[1]); r.bold = True; r.font.size = Pt(10)
                    else: p.add_run(str(frag)).font.size = Pt(10)
            else: p.add_run(str(cd)).font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells): cell.width = col_widths[j]
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
r = p.add_run('HOW-TO DEEP DIVE'); r.bold = True; r.font.size = Pt(8); color_run(r, '8a3ffc')
h1('Systems Integration')
prose('For large-scale enterprises, a capable LLM is not enough. AI transformation requires deep integration into enterprise data and key platforms. Without it, solutions become brittle and hard to maintain. This chapter covers how to surface integration complexity early, write requirements developers can act on, navigate Path to Production, and hand off a solution that lasts.', muted=True)

label('Active across')
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
for ph, sep in [('Engage','→'),('Discover','→'),('Execute','')]:
    colors = {'Engage':'0043a4','Discover':'5c2eb8','Execute':'1a6b32'}
    r = p.add_run(f'[{ph}]'); r.bold = True; r.font.size = Pt(10); color_run(r, colors[ph])
    if sep: p.add_run(f' {sep} ').font.size = Pt(10)

# At a glance
label('At a glance: the six steps')
steps_glance = [
    ('Engage',  '1', 'Tech & Data Assessment',   'Inventory every system and profile data quality before committing to a solution.'),
    ('Engage',  '2', 'Business Process Mapping',  'Document the current-state workflow: systems touched, handoffs, and data flow.'),
    ('Discover','3', 'Workflow Analysis',          "Trace the technical baseline for the workflow you're transforming."),
    ('Discover','4', 'Solution Design',            'Make the build-vs-buy call and write requirements a developer can act on.'),
    ('Execute', '5', 'Experimentation',            'Build and validate the integration against real systems and real data.'),
    ('Execute', '6', 'Scale & Adopt',              'Hand off a documented, owned, maintainable solution.'),
]
phase_colors = {'Engage':'0043a4','Discover':'5c2eb8','Execute':'1a6b32'}
for phase, num, name, what in steps_glance:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f'[{phase} Step {num}]  '); r1.bold = True; r1.font.size = Pt(9); color_run(r1, phase_colors[phase])
    r2 = p.add_run(f'{name} — '); r2.bold = True; r2.font.size = Pt(10)
    r3 = p.add_run(what); r3.font.size = Pt(10); color_run(r3, '525252')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── ENGAGE ────────────────────────────────────────────────────────────────────
phase_tag('Engage', 'Steps 1 & 2 · Tech & Data Assessment · Business Process Mapping')
h2("Know what you're working with before you commit.")
callout('Key Takeaway', "Before you design anything, get an honest picture of two things: what systems exist, and what condition the data is in. Document what's real; don't fix anything yet.")

label('What to do / Key Stakeholder')
add_table(
    ['What to do', 'Key Stakeholder'],
    [
        ['Map the current-state architecture: every system, integration point, and data flow this workflow touches',
         [('bold','System / platform owner'), ' (named, per system)']],
        ['Confirm data quality and access on a real sample, not a description of the data',
         [('bold','Technical & data SME'), ' (BPT/TPT)']],
        ['Document the current business process: handoffs, volume, pain points',
         [('bold','Business process SME')]],
        ['Flag any approval process (architecture or security review) tied to a system already in use',
         [('bold','Platform contact')]],
        ['Escalate a structural blocker (multi-month review, platform migration)',
         [('bold','Domain VP + CIO leadership')]],
    ],
    col_widths=[Cm(10), Cm(7.5)]
)

prose('The Tech & Data Assessment surfaces the constraints that shape every decision downstream. The goal is not to fix anything. It is to document what is real for every system your solution will touch.')
prose('To get started, work with relevant transformation and technical partners to inventory all platforms that are used in your workflow: CRM, ERP, content repositories, and any custom or proprietary systems. Bring in each system owner to confirm: Does an API exist? Is it documented? What are the access requirements? Identify a named contact from each platform team. Platform ownership outside your domain is the most common cause of stagnation.')
prose('Profile data quality on a real sample from each source. Nulls, stale data, and format inconsistencies are constraints to design around. Identify any approval processes that could affect your timeline and start them immediately.')

label('Use the same five questions for every system you touch:')
add_table(
    ['Dimension', 'Question to pressure-test', "Signal it's ready"],
    [
        [('bold','API availability'), 'Does a documented API exist for this system, and can it support your call volume?', 'Standard, documented integration point'],
        [('bold','Access & ownership'), 'Who owns it, and what is the approval process to connect?', 'Named owner, known request path'],
        [('bold','Data freshness'), 'Is the data real-time, or does it update on a delay?', 'Update cadence documented and workable for your use case'],
        [('bold','Data quality'), 'Is the data clean, complete, and usable, not just present?', 'Sample checked; nulls and format issues known'],
        [('bold','Governance'), 'Are there approval processes (architecture or security review) required to connect to this system?', 'Reviewed with platform owner; timeline known'],
    ],
    col_widths=[Cm(3.5), Cm(8), Cm(6)]
)
prose('IBM runs on modern platforms, so the answer to "does it exist" is almost always yes. The real risk is in the other four rows: freshness, ownership, quality, and governance. Surface those in Engage, not mid-build.')

painpoint('"We have the data" is not "the data is usable"',
    'Run a quality check on a real sample: nulls, duplicates, format issues. If 30% of key fields are malformed, that is a constraint to plan around now, not a Sprint 3 surprise. Skipping this check means the build team discovers it instead.')

label('● When to escalate')
bullet([('bold','Small issues '), '(missing API key, undocumented endpoint, format inconsistency): work directly with the platform team.'])
bullet([('bold','Large structural blockers '), '(platform migration, multi-month security review, major infrastructure investment): escalate to the domain VP and CIO leadership before design is finalized.'])
bullet([('bold','Cross-org communication stalled: '), 'escalate to the AIFT team (Radha).'])

painpoint("Skipping technical discovery to 'save time'",
    "When timelines are tight, teams skip the assessment and go straight to design. Integration complexity that surfaces mid-build stops the sprint while the team scrambles to understand systems they should have mapped weeks earlier. The assessment takes two weeks. Skipping it costs two months.")

section_rule()
h3('Cross-Org: Working with CDO: Enterprise Performance Management (EPM)')
prose('If your solution depends on enterprise financial, operational, or planning data, engage the Chief Data Office (CDO) and their Enterprise Performance Management (EPM) platform. EPM is the starting point for understanding what centralized data assets exist and how to access them.')

t_epm = doc.add_table(rows=3, cols=3); t_epm.style = 'Table Grid'
for i, h in enumerate(['Starting point', 'What it covers', 'Link']):
    t_epm.rows[0].cells[i].text = h; t_epm.rows[0].cells[i].paragraphs[0].runs[0].bold = True
epm_data = [
    ('Learn about EPM', 'Understand what EPM is, what data it holds, and whether your use case depends on it',
     'https://w3.ibm.com/epm/askepm', 'w3.ibm.com/epm/askepm'),
    ('Add data to EPM', 'If your solution produces data that should feed back into enterprise reporting, start here to understand the intake process',
     'https://w3.ibm.com/epm/use-epm', 'w3.ibm.com/epm/use-epm'),
]
for ri, (sp, wc, url, lt) in enumerate(epm_data):
    cells = t_epm.rows[ri+1].cells
    cells[0].paragraphs[0].add_run(sp).bold = True
    cells[1].paragraphs[0].add_run(wc).font.size = Pt(10)
    add_hyperlink(cells[2].paragraphs[0], url, lt)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

label('Artifacts')
p = doc.add_paragraph()
p.add_run('Technology & Data Readiness Checklist — ').bold = True; p.runs[-1].font.size = Pt(11)
add_hyperlink(p, 'https://ibm.ent.box.com/file/2181697907530?s=p0abbfft3jumintcy8aw5hewshs2y7fc', 'Example')
p.add_run(' · ').font.size = Pt(11)
add_hyperlink(p, 'https://app.mural.co/t/ibm14/m/ibm14/1775762796754/5e602df33f7af263fd214eeac33feb07e743b699', 'Template')
p2 = doc.add_paragraph()
p2.add_run('Business Process Map(s) — ').bold = True; p2.runs[-1].font.size = Pt(11)
add_hyperlink(p2, 'https://app.mural.co/template/a3fd5e2a-cfdf-4ad6-9887-066521fc6ab1/abc1ccf2-0482-4aa0-8a6c-0cede7949984', 'Example')
p2.add_run(' · ').font.size = Pt(11)
add_hyperlink(p2, 'https://ibm.box.com/s/5c43ream7jq8ezi0boiobmhslhpachwk', 'Template')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

case_study('State University: Discovering the integration landscape',
    "The enrollment team's Tech & Data Assessment revealed three legacy systems with no API documentation, inconsistent data formats across systems, and a mainframe that required batch file exports. The CRM had an API but required security review for external access. The team documented all of this in week one, started the required approval processes immediately, and adjusted their solution design to work with batch exports rather than real-time integration. By surfacing these constraints early, they avoided a mid-build crisis when the development team would have discovered them.")
ready_banner('Ready to move to Discover',
    "You have a documented systems landscape, data quality profiles, access requirements, approval timelines, and a view of where integration complexity is high. You're designing with constraints visible.")

# ── DISCOVER ──────────────────────────────────────────────────────────────────
phase_tag('Discover', 'Steps 3 & 4 · Workflow Analysis · Solution Design')
h2('Map the technical baseline, make the right build decisions, and write requirements developers can act on.')
callout('Key Takeaway', 'Document how data actually moves through the current process, decide build vs. use an existing connector for each capability, and write requirements specific enough that a developer never has to guess.')

label('Who to involve')
add_table(
    ['Who to involve', 'What they contribute', 'When'],
    [
        [('bold','Project architect'), 'Validates technical approach; owns Path to Production navigation; flags required reviews early', 'Start of Discover'],
        [('bold','Product owner'), 'Reviews requirements; confirms they are specific enough for a developer to act on', 'Solution Design'],
        [('bold','Process SMEs'), 'Describe how the workflow runs; confirm the data flow is mapped correctly', 'Workflow Analysis'],
        [('bold','Approval reviewers / Security team'), 'Begin any required architecture or security review processes: do not wait until after requirements are written', 'As soon as scope is known'],
    ],
    col_widths=[Cm(4.5), Cm(10), Cm(3)]
)
prose("Discover has two jobs. In Workflow Analysis (Step 3), you document the current-state architecture for the workflow you're transforming. In Solution Design (Step 4), you make the integration decisions and write the requirements.")

label('Two things to line up before the build starts')
p_p2p = doc.add_paragraph(); p_p2p.paragraph_format.left_indent = Inches(0.25); p_p2p.paragraph_format.space_after = Pt(3)
p_p2p.add_run('Path to Production: ').bold = True; p_p2p.runs[-1].font.size = Pt(11)
p_p2p.add_run("If this will be a production-level solution, IBM's Path to Production process applies (").font.size = Pt(11)
add_hyperlink(p_p2p, 'https://w3.ibm.com/w3publisher/transformation-operations/ai-enablement/build-with-watsonx', 'Build with watsonx')
p_p2p.add_run("), and its requirements are easier to meet when you know them from the start of Discover.").font.size = Pt(11)
bullet([('bold','Resourcing: '), 'Some domains partner with the CIO organization for build resources. If yours does, work with them on prioritization. If resources come from another organization, bring that group into planning the same way.'])

h3('Path to Production process')
p_kt = doc.add_paragraph(); p_kt.paragraph_format.left_indent = Inches(0.25); p_kt.paragraph_format.space_after = Pt(6)
add_hyperlink(p_kt, 'https://w3.ibm.com/w3publisher/transformation-operations/ai-enablement/build-with-watsonx', 'Path to Production')
p_kt.add_run(' is a specific, documented IBM process. Your project architect is a partner in building; the platform team decides whether you can deploy. These are two different relationships with two different timelines. Your platform contact knows which approval processes apply.').font.size = Pt(11)

prose('Path to Production is a named IBM process, run through Build with watsonx, for production-level solutions. It is housed in the CIO/TPT team. The CIO/BPT organization can provide development resources and build capacity, and owns the platform teams that control the production environment.')

p_arb = doc.add_paragraph(); p_arb.paragraph_format.space_after = Pt(4)
p_arb.add_run('Two approval bodies are worth knowing by name. The ').font.size = Pt(11)
p_arb.add_run('Architecture Review Board (ARB)').bold = True; p_arb.runs[-1].font.size = Pt(11)
p_arb.add_run(' is a Transformation & Operations leadership forum, run jointly by BPT and TPT, that decides architecture questions with wide-ranging impact: specifically major platform decisions on core business platforms such as Salesforce (IBM Sales Cloud), SAP, and Adobe. (').font.size = Pt(11)
add_hyperlink(p_arb, 'https://ec.yourlearning.ibm.com/w3/event/10547588', 'ARB information')
p_arb.add_run(') The ').font.size = Pt(11)
p_arb.add_run('Technical Review Board (TRB)').bold = True; p_arb.runs[-1].font.size = Pt(11)
p_arb.add_run(' is part of Path to Production and reviews new AI systems for architecture fit, service readiness, and governance risk. They are different bodies with different scopes: most builds only encounter one, and many encounter neither. Your project architect is the right person to tell you which, if any, applies.').font.size = Pt(11)

label('Which review applies to your build:')
add_table(
    ['Your build involves…', 'ARB?', 'TRB?', 'What to do'],
    [
        ['Core IBM platforms: Salesforce (IBM Sales Cloud), SAP, Adobe', 'Likely yes', 'No', 'Contact your platform owner in week 1 of Discover. Start before requirements are written: takes 4–6 weeks.'],
        ['New internal AI system going to production via Build with watsonx', 'No', 'Likely yes', 'Review requirements at Build with watsonx from the start of Discover. Applies to all IBMers.'],
        ['Internal tool, limited integration scope, no core platform connections', 'Unlikely', 'Unlikely', 'Confirm with your project architect in week 1, not week 8.'],
        ['Not sure', 'Ask your project architect', '↑', 'They know whether a requirement is a real blocker or a standard process with a known path through it.'],
    ],
    col_widths=[Cm(5), Cm(2.5), Cm(2.5), Cm(7.5)]
)

section_rule()
step_num_label(3, 'Workflow Analysis: Document the technical baseline')
prose("In Workflow Analysis, trace the data flow for the workflow you're transforming: where it originates, what systems it passes through, and where it ends. Document every integration point, manual handoff, and place where data quality degrades. This map is the reference for what you're replacing.")
prose('Also identify non-functional requirements that constrain your design: any required architecture reviews, security reviews, data residency rules, and compliance constraints. Some approval processes take 4–6 weeks. Discover that in week eight of a twelve-week build and you\'ve already missed your date.')
painpoint('Underestimating integration complexity',
    'Integration work is underestimated because it\'s invisible until you start building. A system that "has an API" might have one that\'s undocumented, rate-limited, or missing the fields you need. A source that "is accessible" might require a VPN, a service account, and three approval levels. Map integration complexity in Workflow Analysis as a realistic accounting, not an optimistic estimate.')

section_rule()
step_num_label(4, 'Solution Design: Make the build-vs-buy call, then write requirements')
prose("Step 4 has two parts. The first is a strategic decision: for each capability your solution needs, should you build it on IBM's stack, buy an external tool, or combine both? This is not a technical question — it is a question about value, risk, IBM strategy, and what the enterprise is investing in. Get it right before a line of code is written. The second part turns that decision into build-ready requirements that a developer can act on without a clarification call.")
callout('Before you commit',
    "The 3-Lens Framework, the AIFT team's methodology, pressure-tests a technology choice from three angles before you jump to build-vs-buy: does it solve the real business problem, can you actually build and run it, and does the decision still make sense as the enterprise and market evolve? A strong recommendation is a plausible yes to all three, not just the one that's easiest to answer.")

h3('Business fit: does this solve the actual business problem?')
add_table(['Dimension','Question to pressure-test','Signal of fit'],
    [[('bold','Value'),'What pain point is this actually solving, for whom?','Tied to a named metric: productivity, revenue, cost, time savings, risk'],
     [('bold','UX fit'),'Is this less friction than the current workflow?','Removes steps or context-switching people already do'],
     [('bold','Adoption'),'Will people actually use it, and keep using it?','Users trust the output and have a reason to opt in'],
     [('bold','ROI'),'Do the benefits justify the cost and change effort?','You are seeing the dollar benefit from the solution']],
    col_widths=[Cm(3), Cm(8), Cm(6.5)])
prose('AIFT example: developing enterprise data agents for the sales domain to improve access to performance insights data.', small=True, muted=True)
prose('Start with the workflow, not the tool.')

h3('Technical fit: can we make it reliably work and scale?')
add_table(['Dimension','Question to pressure-test','Signal of fit'],
    [[('bold','Architecture'),'How does it connect to and fit into our systems and data?','Documented, standard integration points'],
     [('bold','Security / Governance'),'Who has access, and can we prove it later?','SSO, audit logs, CISO requirements met'],
     [('bold','Scalability'),'Does it hold up under real usage, not a demo?','Load-tested against actual peak volumes'],
     [('bold','Reliability'),'What happens when something fails?','Clear error handling, alerting, uptime targets'],
     [('bold','Data readiness'),'Is the underlying data actually usable?','Clean, governed, well-defined sources']],
    col_widths=[Cm(3.5), Cm(8), Cm(6)])
prose('AIFT example: can we leverage dbt transformation to support an AI agent\'s ability to work with IBM enterprise data, enabling consumption via natural language by an IBMer?', small=True, muted=True)
prose('Separate demo feasibility from production feasibility. Better technology or a better model is not a better solution if it cannot integrate, scale, or be governed.')

h3('Strategic fit: does this still make sense as things evolve?')
add_table(['Dimension','Question to pressure-test','Signal of fit'],
    [[('bold','Aligns with IBM strategy'),"Does this reinforce where we're already placing our bets, or compete with them for budget?",'Maps cleanly to an existing investment priority'],
     [('bold','Future flexibility'),'Can we change direction later without starting over?','Modular, standards-based, portable outputs'],
     [('bold','Lock-in risk'),'How hard would it be to exit from a vendor or solution if needed?','Clear exit path with bounded switching cost'],
     [('bold','Capability building'),'What internal skills do we want to own versus rent?','Team builds real fluency, not just usage habits'],
     [('bold','Cost trajectory'),'How do costs scale as usage grows?','Predictable, linear, or improving unit economics']],
    col_widths=[Cm(3.5), Cm(8), Cm(6)])
prose('AIFT example: should we be thinking about buying third-party LLM front-end solutions (e.g. Gemini Enterprise, Claude Enterprise) to support IBMers with day-to-day productivity?', small=True, muted=True)
prose('Distinguish what is optimal today from what is durable tomorrow.')

label('Synthesizing your assessment: make the trade-offs visible')
add_table(['Business Fit','Technical Fit','Strategic Fit','Implication'],
    [['High','High','High','Scale with confidence'],
     ['High','Low','Medium','Build a proof of concept, then de-risk the architecture'],
     ['Low','High','High','Do not over-engineer a weak use case'],
     ['High','High','Low','Proceed, but identify exit options and set roadmap guardrails']],
    col_widths=[Cm(2.5), Cm(2.5), Cm(2.5), Cm(10)])
prose('A strong recommendation shows what you are optimizing for and what risk you are accepting, not just a single score. With the assessment done, the build-vs-buy call for each integration point follows directly.')
prose("Once you've synthesized your 3-lens assessment, the decision for each integration point typically resolves to one of four categories.")

label('⚠ The trap')
bullet([('bold','Defaulting to buy for speed '), 'without checking lock-in, data exposure, and total cost.'])
bullet([('bold','Defaulting to build '), 'without checking whether a connector already exists.'])
bullet(['Not documenting the reasoning: you will need to defend the decision when questions arise during the build.'])

label('The four options:')
add_table(['Option','When to use it','IBM examples'],
    [[('bold',"Build: use IBM's stack"),"IBM has the right capability, or it is being built. IBM acquisitions count as the IBM stack. This is also where Client Zero proof points come from.",'watsonx, Orchestrate connector catalog, IBM acquisition tooling'],
     [('bold','Buy: go external'),'IBM does not have the right fit yet, and mature tools exist with proven AI functionality.','Quote-to-Cash, Source-to-Pay, legal review: most have purpose-built marketplace tools'],
     [('bold','Hybrid: combine'),'IBM capabilities plus external tools for the gaps. This is also how ecosystem partnerships start.','IBM model + external workflow tooling'],
     [('bold','Sequence: short-term / long-term'),'Run a third-party tool short-term (a year or until an IBM product exists). Plan the transition from day one.','Use vendor tool now; migrate to IBM stack when ready']],
    col_widths=[Cm(4.5), Cm(8.5), Cm(4.5)])

# Fast Track callout
label('Experimenting with an external tool? Use Fast Track.')
p_ft = doc.add_paragraph(); p_ft.paragraph_format.left_indent = Inches(0.25); p_ft.paragraph_format.space_after = Pt(4)
p_ft.add_run("If your team wants to pilot a third-party AI tool before committing to buy, ").font.size = Pt(11)
p_ft.add_run("Fast Track").bold = True; p_ft.runs[-1].font.size = Pt(11)
p_ft.add_run(" is IBM's process for it. It enables teams to quickly evaluate external AI tools for enterprise risk and value — with coordination from CIO, CISO, Responsible Tech, Legal, and Procurement — in a 90-day, up to 500-license pilot scoped to publicly available IBM data and no system integrations.").font.size = Pt(11)
p_ft2 = doc.add_paragraph(); p_ft2.paragraph_format.left_indent = Inches(0.25); p_ft2.paragraph_format.space_after = Pt(8)
p_ft2.add_run('Learn more: ').font.size = Pt(11)
add_hyperlink(p_ft2, 'https://w3.ibm.com/w3publisher/ai-first-transformation/key-initiatives/fast-track', 'Fast Track overview')
p_ft2.add_run('  ·  ').font.size = Pt(11)
add_hyperlink(p_ft2, 'https://forms.monday.com/forms/b51e91524820fb0a0acb4c7636302f2f?r=use1', 'Submit a request')

prose("Once you've made the build-vs-buy decision, write the requirements. Test them by asking whether a developer can start work without a clarification call. Build-ready requirements include:")
bullet([('bold','Persona: '), 'Who is using this?'])
bullet([('bold','Data objects: '), 'What specific fields, format, and source system?'])
bullet([('bold','Systems involved: '), 'Which API endpoint, authentication method, and rate limits?'])
bullet([('bold','Acceptance criteria: '), 'What does done look like?'])
bullet([('bold','Fallback path: '), 'What happens when the integration fails?'])
prose('Any gap in those fields means the developer stops and asks. That is when timelines slip.')

painpoint("Requirements 'too conceptual to build from'",
    'Business teams write requirements like "integrate with CRM to pull customer data", which sounds clear but leaves a developer asking: which CRM? Which endpoint? Which fields? What format? What happens if the API is down? Requirements are build-ready when none of those questions remain. If your document is under three pages, it is probably too conceptual.')

label('Before you move to Execute')
bullet([('bold','Test the requirements '), 'by giving them to a developer. If they need to ask clarifying questions, add more specificity.'])
bullet([('bold','Document the tech stack decision '), "with pros, cons, and tradeoffs. You'll need to defend it when questions arise later."])

section_rule()
h3('Best Practice: Working with BPT-managed platforms')
prose('Once you know which platforms your solution depends on, check whether they are BPT-managed. If they are, the PML (Product Management Leader) is your point of contact, not a generic IT ticket.')
t_pml = doc.add_table(rows=5, cols=3); t_pml.style = 'Table Grid'
for i, h in enumerate(['Step','Action','How / Who']):
    t_pml.rows[0].cells[i].text = h; t_pml.rows[0].cells[i].paragraphs[0].runs[0].bold = True
pml_rows = [
    ('1','Identify the PML for each BPT-managed platform', ('Use the ', 'https://ibm-my.sharepoint.com/:p:/p/emilia_ponechalova_sk/IQBle-NYPLimTJhnDgXm8CXVAY20e5Ti8QMD2pUwHxNAx8c?e=z6HaS2', 'PML Chart'), '. PMLs are the named point of contact within each platform tower.'),
    ('2',"Understand the platform's request and prioritization process", None, 'Work with the PML: they own how requests are submitted and ranked within the platform org. Timelines vary by platform.'),
    ('3','Validate technical integration details', None, 'The PML can confirm API patterns, access requirements, and whether your use case requires customization work from their team.'),
    ('4','Use the PML as your escalation path', None, 'Platform leaders coordinate with business unit focals on prioritization. The PML moves your request forward within the platform org.'),
]
for ri, row in enumerate(pml_rows):
    cells = t_pml.rows[ri+1].cells
    cells[0].paragraphs[0].add_run(row[0]).bold = True
    cells[1].paragraphs[0].add_run(row[1]).font.size = Pt(10)
    p_how = cells[2].paragraphs[0]
    if row[2]:
        p_how.add_run(row[2][0]).font.size = Pt(10)
        add_hyperlink(p_how, row[2][1], row[2][2])
        p_how.add_run(row[3]).font.size = Pt(10)
    else:
        p_how.add_run(row[3]).font.size = Pt(10)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

label('Deliverable')
p_bvb = doc.add_paragraph()
p_bvb.add_run('Build vs Buy — ').bold = True; p_bvb.runs[-1].font.size = Pt(11)
add_hyperlink(p_bvb, 'https://ibm.sharepoint.com/:p:/r/sites/AI-FirstTransformation_DEPT/_layouts/15/Doc.aspx?sourcedoc=%7BB802C135-3D51-45A7-AF1C-E33A9CFF945E%7D&file=Build%20vs%20Buy%20Example.pptx&action=edit&mobileredirect=true', 'Example')
p_bvb.add_run(' · ').font.size = Pt(11)
add_hyperlink(p_bvb, 'https://ibm.sharepoint.com/:p:/r/sites/AI-FirstTransformation_DEPT/_layouts/15/Doc.aspx?sourcedoc=%7BABAEDCF6-1BA8-4419-8AD7-74364B462947%7D&file=Build_vs_Buy_Comparison_Template.pptx&action=edit&mobileredirect=true', 'Template')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

case_study('State University: Build-vs-buy decision',
    "For the document verification workflow, the team faced three integration points: CRM data pull, document storage access, and ERP update. They used the matrix: CRM integration was high-value, low-complexity (existing Salesforce connector), buy and customize. Document storage was high-value, high-complexity (custom mainframe interface), partner with IT to build. ERP update was low-value, high-complexity, descoped to post-MVP. The team documented each decision with reasoning, which proved critical when the CFO questioned why they weren't integrating with ERP in the first release.")
ready_banner('Ready to move to Execute',
    "You have build-ready requirements, a documented build-vs-buy decision for each integration point, any required approval reviews in progress, and a tech stack decision with documented tradeoffs. The build team can start.")

# ── EXECUTE ───────────────────────────────────────────────────────────────────
phase_tag('Execute', 'Steps 5 & 6 · Experimentation · Scale & Adopt')
h2('Build, integrate, instrument, and hand off cleanly.')
callout('Key Takeaway', 'Connect to real systems from Sprint 1, test every way the integration can fail, instrument it so problems surface before users report them, and hand off to one named person before the build team steps back.')

label('Who to involve')
add_table(['Who to involve','What they contribute','When'],
    [[('bold','Project architect'),"Bridges to the platform team; keeps approval reviews moving alongside the build so they don't land on launch day",'Throughout Execute'],
     [('bold','Developer lead'),'Builds and integrates on real data from Sprint 1; implements error handling and instrumentation from day one','Sprint 1 through launch'],
     [('bold','Product owner'),'Confirms UAT is conducted on real integrations, not mocked data; signs off on production readiness','UAT and launch'],
     [('bold','Platform team'),'Controls the production environment and deployment approval; separate conversation from the build team','Path to Production reviews'],
     [('bold','Named technical owner'),'Shadows the build team; learns the solution before handoff; confirms they can troubleshoot independently','Final sprint through handoff']],
    col_widths=[Cm(4), Cm(10), Cm(3.5)])
prose("Execute is where the requirements from Discover get tested against real systems, real data, and real failure modes. The goal is not a working demo; it's a solution that runs in production, handles errors, and can be maintained by someone other than the person who built it.")

step_num_label(5, 'Experimentation: Build and validate on real data')
prose("Connect to real systems and real data now. The quality issues you find (missing fields, format inconsistencies, unexpected nulls) aren't surprises; they're confirmation of what you documented in Engage, now needing to be handled in code: error handling on every call, logging on every step, and pipelines built for production volume from day one.")

painpoint('Data quality worse than expected',
    'The quality profile from Engage was on a sample. The full dataset will have edge cases the sample missed: nulls, mismatched IDs, unexpected date formats. This is expected. Build your pipeline to handle bad data: log it, flag it, and keep processing. A pipeline that crashes on the first bad record is not production-ready.')
prose('Instrument everything. Log what data came in, what transformations ran, what went out, and whether it succeeded. When something breaks in production, you need a traceable record. Integrations that are easy to maintain have every step logged and every error visible.')
painpoint("One-off integrations that can't scale",
    "Teams build integrations that work for MVP volumes but break at production scale. A script handling 100 records a day fails at 10,000. An API call that passes in testing fails in production due to rate limits. Build for scale from the start: batch processing for high volumes, rate limiting, retry logic. An integration that fails at scale is worse than a known gap.")

section_rule()
step_num_label(6, 'Scale & Adopt: Hand off a maintainable solution')
prose("The handoff to production is where most integrations fail. Not because the code doesn't work, but because no one knows how to maintain it after the build team leaves. A successful handoff requires complete technical documentation, a named owner who can actually run and maintain the solution, and a clear escalation path when something breaks. If any item in the checklist below is missing, the solution will degrade quietly until it stops working entirely.")

label('Technical handoff checklist')
label('Code & Config')
bullet([('bold','Code repository with README '), '— documented, with test cases and setup instructions'])
bullet([('bold','Secrets & API keys '), '— with rotation instructions and access procedures'])
bullet([('bold','Environment configuration '), '— dev, test, prod settings and deployment process'])
label('Architecture & Data')
bullet([('bold','Architecture diagram '), '— all systems, integration points, data flows'])
bullet([('bold','Data lineage documentation '), '— where each field comes from and how it\'s transformed'])
bullet([('bold','API documentation '), '— endpoints, authentication, rate limits, error codes'])
label('Operations & Support')
bullet([('bold','Runbook for common failures '), '— what to do when API fails, data is bad, system is down'])
bullet([('bold','Monitoring & alerting setup '), '— dashboards, alerts, escalation procedures'])
bullet([('bold','Named technical owner '), '— person who can run, maintain, and troubleshoot'])
bullet([('bold','Escalation path '), "— who to contact when owner can't resolve issue"])
prose('⚠ If any item is missing, the solution will degrade after handoff. Complete the checklist before the build team leaves.')

prose("The named owner is not just a name on paper. They need to have been trained on the solution, run through failure scenarios with the build team, and confirmed they can troubleshoot without escalating back to you. If they can't, the handoff isn't done.")
painpoint('No named technical owner after handoff',
    'The most common handoff failure is transferring the solution to "the IT team" or "the platform team" without a specific person who owns it. When something breaks, no one knows who\'s responsible for fixing it. The issue gets escalated, bounced between teams, and eventually the solution stops working. Avoid this by naming a specific owner before the handoff, confirming they have the skills and capacity to maintain it, and documenting their name in every piece of handoff documentation. If you can\'t name a person, the handoff isn\'t ready.')

case_study('State University: Clean handoff to IT',
    "Before the enrollment transformation team stepped back, they completed a full technical handoff to the IT integration specialist. The handoff included: documented code repository with test cases, architecture diagram showing all three system connections, API keys with rotation procedures, data lineage documentation, and a runbook for the five most common failure scenarios. The specialist shadowed the build team for two weeks, ran through each failure scenario, and confirmed she could troubleshoot independently. Six months later, when the CRM API changed and broke the integration, the specialist diagnosed and fixed it in under two hours using the runbook, without needing to escalate to the original build team.")
ready_banner('Ready to move to the next workflow',
    "You have a complete technical handoff package, a named owner who can maintain the solution, monitoring and alerting in place, and documented limitations and next steps. The solution will survive the transition to production and continue working after the build team leaves.")

out = '/Users/claireliu/Desktop/systems-integration-chapter_V7.docx'
doc.save(out)
print(f'Saved: {out}')
