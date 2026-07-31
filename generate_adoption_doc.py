"""
Generate adoption-change.docx from the Adoption & Change deep dive narrative.
Run: python3 generate_adoption_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Margins ───────────────────────────────────────────────────────────────────
s = doc.sections[0]
s.page_width  = Inches(8.5)
s.page_height = Inches(11)
s.left_margin = s.right_margin = Inches(1)
s.top_margin  = s.bottom_margin = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x16, 0x16)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x16, 0x16)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x52, 0x52, 0x52)
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def callout(label, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"⚠ {label}: {title}\n")
    r.bold = True
    p.add_run(text)
    return p

def note(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    return p

def case_study(context, *paras):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(f"Case Study — {context}\n")
    r.bold = True
    r.font.color.rgb = RGBColor(0x5c, 0x2e, 0xb8)
    for para in paras:
        p.add_run(para + "\n")
    return p

def table_3col(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = val
    return t

def table_ncol(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = val
    return t

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Adoption & Change", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph(
    "How-to Deep Dive — Active across Engage → Discover → Execute\n\n"
    "Deploying technology is the easy part. Getting people to change how they work is where most "
    "AI transformations stall. This chapter covers how to bring users along from day one, design "
    "communications that actually land, run training that sticks, and sustain adoption after the "
    "build team is gone."
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════════════════════════════════
# AT A GLANCE
# ══════════════════════════════════════════════════════════════════════════════
h1("At a Glance — Where Adoption & Change Lives")
table_ncol(
    ["Step", "Phase", "What to do for Adoption & Change"],
    [
        ("Step 1 · Tech & Data Assessment", "Engage",
         "Map who is affected, what changes for them, what they lose, and what concerns they are likely to have."),
        ("Step 2 · Business Process Mapping", "Engage",
         "Re-envision what people will actually do in the new workflow — not just the technology — and get the executive sponsor on record."),
        ("Step 3 · Workflow Analysis", "Discover",
         "Recruit real UAT users (frontline workers, not managers) before design is finished. Your testers become your first advocates."),
        ("Step 4 · Solution Design", "Discover",
         "Plan communications for the whole project and design training before the build, including what to do when the AI is wrong."),
        ("Step 5 · Experimentation", "Execute",
         "Start small with real work, drive usage, and treat resistance as data about what to fix."),
        ("Step 6 · Scale & Adopt", "Execute",
         "Launch with training done, monitor adoption weekly, intervene early, and retire the old process."),
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# FRAMING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Transformation Is Three Things, Not One")

body(
    "Every AI workflow transformation involves technology, people, and process. Most projects spend "
    "90% of their energy on technology, forget that process needs to be re-envisioned alongside it, "
    "and treat people as an afterthought to be trained at the end. Gartner found that half of "
    "enterprise AI project failures are change management failures — not technical ones. The model "
    "worked. The people did not adopt it."
)

table_3col(
    ["Technology", "People", "Process"],
    [("The AI solution, integrations, infrastructure. Necessary but not sufficient.",
      "The humans whose daily work changes. Their awareness, willingness, skill, and habit determine whether the technology actually gets used.",
      "The end-to-end workflow re-envisioned around what AI makes possible. Not bolting AI onto the existing process, but rethinking how the work gets done.")]
)

body(
    "Change management is not a separate workstream that runs after the build. It runs throughout "
    "the project. In Engage, it is mapping who is affected and getting leaders aligned. In Discovery, "
    "it is involving real end users in design and starting communications. In Execute, it is running "
    "training before launch (not after), using UAT as an enablement activity, and monitoring adoption "
    "closely enough to step in before a plateau becomes a loss."
)
body(
    "The domain owns its own change management. The AIFT team provides the framework and support. "
    "But the people leading the domain know their colleagues, know the culture, and know the specific "
    "risks and concerns that will drive resistance. That local knowledge is the most important input "
    "to a change plan."
)

# ══════════════════════════════════════════════════════════════════════════════
# ENGAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Engage Phase — Steps 1 & 2")
h2("Map who is affected, re-envision the process, and get leaders on record before design starts.")

note("Key takeaway",
     "Map who is affected and what concerns they have, re-envision what people will actually do "
     "in the new workflow, and get the sponsor on record. All of it before design starts — because "
     "after the build, those decisions are already made.")

note("Who to involve",
     "Domain transformation lead · Executive sponsor · Change owner · Representative end users")

body(
    "Change management that starts at launch is too late. By the time the solution is built, the "
    "decisions that most affect user experience — what the new workflow looks like, who does what "
    "differently, what tasks disappear — have already been made without user input. Starting in "
    "Engage means those decisions are informed by the people who will live with them."
)

h3("What to do in Engage for Adoption & Change")
numbered("Map affected personas.",
         "Who does this workflow today, what will change for each, and what concerns they are likely to have. Include frontline users, managers, subject-matter experts, and adjacent teams.")
numbered("Name the change owner.",
         "This person is responsible for the people-and-process side of the transformation. They are not the domain transformation lead and not the product owner. They own the change plan, the communications, and the adoption tracking.")
numbered("Re-envision the future-state workflow with input from frontline users.",
         "Document what people will do differently, not just what the AI will do.")
numbered("Get the executive sponsor to communicate the vision.",
         "Why this change is happening, what it means for the team, and what they expect from people going forward.")
numbered("Name at least two champions.",
         "Respected peers (not managers) who will influence adoption from within the team. Champions need early access to the solution and a specific role: they give feedback during design and carry the message to colleagues during rollout.")

h3("Step 1 · Change Impact Assessment: Understand Who Is Affected and How")
body(
    "The first change management activity is a high-level read of who is affected. This is not a "
    "detailed analysis. It is enough to surface the main personas whose roles will change, and a "
    "rough sense of how those people are likely to react."
)
body(
    "For each affected persona, ask three questions: What specifically changes in their day-to-day "
    "work? What do they lose (autonomy, familiar steps, perceived expertise)? And what concerns might "
    "they have (being replaced, being evaluated differently, having to learn something new)? That last "
    "question is the most important and the one most teams skip."
)

table_ncol(
    ["Persona", "What changes", "What they might lose", "Likely reaction"],
    [
        ("Frontline user", "Core task is now AI-assisted or automated",
         "Mastery of a familiar process; sense of contribution",
         "Skepticism, avoidance, or active workaround if not brought in early"),
        ("Manager / team lead", "How they supervise and what they measure changes",
         "Visibility into work; familiar management rhythms",
         "Neutral to positive if outcomes improve; resistant if they lose oversight"),
        ("Subject-matter expert", "Their specialized knowledge is now encoded in the model",
         "Status as the person who knows; request volume",
         "Often the most resistant — or, if brought in as a champion, the most influential"),
        ("Adjacent team", "Inputs or outputs they depend on change format or timing",
         "Predictable handoffs; established communication patterns",
         "Confused or frustrated if not notified early; easy to bring on board if informed in advance"),
    ]
)

callout("Pain point", "Not addressing the job security question",
        "Across IBM domains, the single most consistent source of adoption resistance is the concern "
        "that the AI is going to replace the person using it. If that question is not named and answered "
        "directly, it will run as an undercurrent through everything else. Tell people what the AI handles, "
        "what it does not handle, and what their role looks like going forward. Be specific. "
        "Vague reassurances only deepen the concern.")

h3("Step 2 · Process Re-envisioning: Redesign the Workflow, Not Just the Technology")
body(
    "The most common structural failure in AI workflow transformation is designing the AI solution "
    "and then asking 'what will people do with this?' That question needs to come first, not last. "
    "In Business Process Mapping (Step 2), the goal is not just to map the current process. It is to "
    "re-envision the end-to-end future-state workflow: the specific steps people will take, what AI "
    "handles, what humans handle, and what the handoffs look like."
)

callout("Pain point", "Late scramble: 'what will they actually do with this?'",
        "Teams design the technology in detail and the new workflow only vaguely. By the time the "
        "solution is close to finished, no one has a clear answer to 'what will the analyst actually "
        "do different on Monday morning?' The future-state workflow should be documented in Business "
        "Process Mapping (Step 2), not discovered in UAT.")

body(
    "Get leaders on record in Engage. Leadership alignment is not a one-time announcement — it is "
    "a sustained behavior. The most powerful signal users receive about whether a change is real is "
    "whether their manager takes it as a real priority. Get the executive sponsor on record with a "
    "clear statement of why this transformation is happening, what it means for the team, and what "
    "they expect from people."
)

case_study(
    "MCC Lesseps: framing the change and building trust from day one",
    "The MCC Lesseps team framed the case for change in concrete workflow terms: manual, fragmented "
    "pre-event work was creating bottlenecks that constrained speed and consistency. Not 'we need to "
    "adopt AI,' but a specific problem the team already felt. Leaders modeled the AI-supported workflow "
    "in their own work rather than delegating the change downward, and the AI was positioned as "
    "'a fast collaborator, not an authority.' That framing set the trust expectation early: the tool "
    "proposes, people decide. Users came into UAT expecting to exercise judgment, not to be replaced."
)

note("Artifact", "Change Management Plan — Started in Engage with personas, change impact, change owner, champions, and the executive sponsor communication. Built out through Discovery and Execute as details are confirmed.")

# ══════════════════════════════════════════════════════════════════════════════
# DISCOVER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Discover Phase — Steps 3 & 4")
h2("Involve real users in design. Plan communications and training before the build starts.")

note("Key takeaway",
     "Three decisions before the build starts: who tests it (real end users, 5–8 of them, not "
     "managers), how you communicate at each stage, and how people learn it — including what to "
     "do when the AI is wrong.")

note("Who to involve",
     "Change owner · Representative end users · Domain transformation lead · Champions · Communications lead (if available)")

body(
    "Discovery is where the change management plan gets specific. You have a picture of who is "
    "affected from Engage. Now you use that picture to make three decisions: who participates in "
    "UAT, what the communications plan looks like across the whole project, and how training will "
    "be delivered. All three need to be decided before the build starts."
)

h3("What to Design in Discovery for Adoption & Change")
numbered("Recruit UAT participants.",
         "Frontline workers who do the actual work, not managers. Aim for 5–8 people who represent "
         "the range of users: different experience levels, different use patterns, different levels "
         "of enthusiasm about change.")
numbered("Write the communications plan.",
         "For each phase (pre-build, pre-launch, launch, post-launch), name the message, the audience, "
         "the channel, and who sends it. The executive sponsor sends the first message. A manager or "
         "champion sends the training reminder. Peer success stories come after launch.")
numbered("Define the 'what is in it for me' message per persona.",
         "One message for frontline users (time back, less drudgery, more interesting work). A "
         "different message for managers (better visibility, faster outcomes). A different message "
         "for subject-matter experts (their knowledge preserved and scaled).")
numbered("Design the training approach before the build starts.",
         "Decide: is this workflow-embedded (a job aid on the screen while they work), a short "
         "guided walkthrough before first use, or a peer demo from a champion? Training should "
         "cover the happy path, what to do when the AI is wrong, and where to get help.")
numbered("Build a feedback channel into the solution itself.",
         "Not as a separate survey. Users who encounter problems need an immediate way to flag them. "
         "That feedback drives Sprint 2 and 3 improvements that turn an imperfect MVP into something "
         "people want to use.")

h3("Step 3 · Recruit Real UAT Users Before Design Is Finished")
body(
    "UAT is the single best adoption lever available during the build. The people who test the "
    "solution before launch become its first advocates after launch — if you involve the right people. "
    "The right people are frontline workers who do the actual work, not the managers who supervise it."
)
body(
    "Recruit UAT participants in Discovery, not the week before testing starts. Give them enough "
    "context about what the solution is designed to do that their feedback is informed. When they find "
    "problems (and they will), close the feedback loop fast. An open feedback loop that users can "
    "see builds more trust than a polished demo."
)

callout("Pain point", "Running UAT with managers instead of end users",
        "The most consistent UAT mistake is recruiting the people who are easy to schedule — managers, "
        "leads, power users — instead of the people who will use the tool every day. A manager who sees "
        "a 40% efficiency improvement in aggregate may approve the solution with enthusiasm. A frontline "
        "worker who now has to handle the 5% of cases the AI gets wrong, without any guidance on how "
        "to do that, will avoid the tool. Test with the people who will live with the solution.")

h3("Step 4 · Plan Communications: Right Message, Right Time, Right Channel")
body(
    "Communications for an AI transformation are not a launch announcement. They are a sustained "
    "conversation that starts in Engage and continues after the solution goes live."
)

table_3col(
    ["Phase", "Message", "Channel"],
    [
        ("Engage / Discovery",
         "Why this is happening. What problem we are solving. Why now. What role you will have in shaping it.",
         "Team meeting, executive message"),
        ("Build (before launch)",
         "What is changing and what it means for you. Specific changes to the workflow. What you will do differently. What AI handles. What your role still owns.",
         "Team briefing, job aid, FAQ"),
        ("Launch and after",
         "How to use it and where to get help. Hands-on guidance, edge case handling, feedback channel. Regular updates on how adoption is going.",
         "Training, office hours, Slack, manager reinforcement"),
    ]
)

body(
    "The most important communication is the 'what is in it for me' message for each persona. "
    "A counselor who spends 60% of their time on manual document checks wants to hear that they "
    "will spend that time on higher-value advising instead. A subject-matter expert wants to hear "
    "that their knowledge is being preserved, not replaced."
)

callout("Pain point", "Announcing the change once and assuming people heard it",
        "A single launch email does not change behavior. People need to encounter the message "
        "multiple times, in multiple formats, from multiple voices: their manager, a peer champion, "
        "a brief in a team meeting, a follow-up from the executive sponsor. And the message needs "
        "to be specific: 'starting Monday, when an order fails validation, the system will flag it "
        "automatically and send you the exception with a reason code' lands. "
        "'We are implementing an AI tool' does not.")

h3("Step 4 · Design Training Before the Build Starts")
body(
    "Training that happens after launch misses the window when people are most open to it. The best "
    "moment to train someone on a new tool is before they use it for the first time. Training designed "
    "after the solution is built is often rushed, covers the happy path, and misses the edge cases "
    "users will encounter in their first week."
)
body(
    "Training for AI tools is not the same as training for software. People need to understand not "
    "just how to use the tool but when to trust it, when to check its output, and what to do when it "
    "is wrong. That last point — how to handle AI errors — is the training content most often missing "
    "from rollouts, and it is the gap that drives avoidance."
)

case_study(
    "MCC Lesseps: tailored messaging and learning in the flow of work",
    "The MCC Lesseps team ran persona-specific engagement rather than one broadcast: event owners, "
    "managers, and leadership each got messaging tailored to what would change for them. Training was "
    "embedded in the workflow itself (short guidance and examples at the point of use) instead of "
    "formal training courses, so people learned the tool while doing real work. And UAT was run on "
    "real events, not sandbox scenarios, which meant testers built competence on the actual work "
    "they would be doing after launch."
)

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTE — STEP 5
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Execute Phase — Step 5: Experimentation")
h2("Start small with real work. Launch with training already done. Monitor adoption and intervene early.")

note("Key takeaway",
     "Start with one area of real work and expand as trust builds. Watch the adoption numbers "
     "weekly, treat resistance as data about what to fix, and retire the old process so there is "
     "no path back to the status quo.")

note("Who to involve",
     "Change owner · Domain transformation lead · Champions · Manager / team lead · Executive sponsor")

body(
    "The most powerful adoption technique is the one that sounds the most obvious: have real people "
    "use the actual tool for real work, starting small. Not a demo. Not a sandbox. A real task, "
    "with a real outcome, inside the actual workflow. What happens at that first moment of real use "
    "determines whether the person becomes an advocate or an avoider."
)

h3("How to Run the Rollout and Sustain Adoption")
numbered("Choose your rollout approach.",
         "Two options: phased (one team first, expand after 2–4 weeks of stable adoption) or all-at-once "
         "(full user group on day one). Phased rollouts give you a contained feedback loop. "
         "All-at-once works when training is solid, the change is low-risk, and the user group is small.")
numbered("Complete training before launch day, not on it.",
         "Workflow-embedded training (a job aid visible in the tool, a 15-minute peer walkthrough) "
         "is more effective than a group session the week after launch.")
numbered("Retire the old process explicitly on launch day.",
         "Communicate what is being turned off, when, and what to do if a case requires a different "
         "approach. Vague exceptions become general exceptions.")
numbered("Track adoption metrics weekly for the first eight weeks.",
         "Assign the change owner to review the numbers every week and flag any plateau or drop "
         "within 48 hours.")
numbered("Respond to resistance with investigation, not pressure.",
         "When a user is not adopting, find out why before deciding what to do. A champion "
         "conversation works better than a manager mandate for most resistance patterns.")
numbered("Close the feedback loop visibly.",
         "When users flag a problem and see it fixed in the next sprint, they trust the channel. "
         "When they flag problems and hear nothing, they stop flagging and stop using the tool.")
numbered("Celebrate the first real wins with the team that created them.",
         "A counselor who reduced verification time from 14 hours to 2 hours is a more powerful "
         "advocate than any communication from leadership. Share that story, with specifics.")

body(
    "Start with a narrow scope and one or two willing participants. The goal is to build a proof "
    "point that a real person can point to. One concrete result from a real colleague is worth more "
    "than a hundred slides. A hesitant manager who sees a specific colleague succeed is far more "
    "likely to try it themselves than one who has only heard an abstract case."
)

case_study(
    "R2A Finance: start small to drive the mindset shift",
    "The R2A team had a finance manager who was skeptical about the AI analysis tool. Rather than "
    "arguing the case, the team asked him to pick just one of the seven analysis areas his team "
    "handled and use the tool for real work in that area only. No other tools allowed for that "
    "analysis. The constraint was deliberate: it forced real usage rather than hedged comparison.",
    "The next day, the manager came back and asked to run the tool across four areas that month, "
    "and all seven the month after. He went from resistant to advocate in 24 hours — not because "
    "the tool was perfect, but because he did real work with it and saw a real result."
)

h3("Resistance Patterns and How to Respond")
body("Resistance is data, not a problem to overcome. When users avoid the tool, investigate before intervening.")

table_ncol(
    ["Signal", "Pattern", "Response"],
    [
        ("Avoidance", "Trust gap",
         "User does not trust the AI output enough to act on it. Add explainability, share accuracy data, let champions demo real results."),
        ("Workaround", "Workflow friction",
         "The new process is harder than the old one for a specific case. Investigate the specific case, fix the friction in the next sprint, or document the exception handling path."),
        ("Minimal use", "Training gap",
         "User knows the tool exists but does not feel confident enough to use it. Targeted one-on-one walkthrough from a champion — not another group training session."),
        ("Vocal pushback", "Concern not addressed",
         "User has an unaddressed concern about what the change means for their role. A direct conversation: 'What specifically do you think will change?' Answer it directly."),
    ]
)

body(
    "The manager's role during Experimentation is not to monitor compliance. It is to visibly use "
    "the tool themselves and reinforce the new workflow in how they talk about the work. If the "
    "manager keeps referring to metrics from the old process, the team will keep working the old way."
)

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTE — STEP 6
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Execute Phase — Step 6: Scale & Adopt")
h2("Launch with training done, monitor weekly, retire the old way.")

body(
    "By the time the solution is ready to launch to the full user group, three things need to "
    "already be in place: training completed (not scheduled), the old process retired (not just "
    "'no longer required'), and the adoption metrics defined and tracked. Launching without "
    "training is the most consistent cause of adoption plateau. Leaving the old process available "
    "is the most consistent cause of reversion."
)

body(
    "Training before launch, not after. The window when people are most open to learning is before "
    "they use the tool for the first time in a real context. Training should happen in the week "
    "before launch, be specific to the actual workflow, and include what to do when the AI is wrong."
)

body(
    "Retire the old process explicitly. As long as the old way remains available, there will be "
    "a subset of users who use it — for 'complex cases,' for 'times when the AI isn't quite right,' "
    "or out of habit. If the old process needs to stay available for some cases, document exactly "
    "which cases and what triggers them. Vague exceptions become general exceptions."
)

callout("Pain point", "Declaring victory at launch",
        "The most common adoption failure is treating launch as the end of the change effort. "
        "Adoption at the end of week one is a lagging indicator of training quality, not a leading "
        "indicator of sustained use. Real adoption takes 6–8 weeks to establish as habit. The teams "
        "that sustain high adoption track usage weekly, investigate every plateau right away, and "
        "stay close to users for the full post-launch period — not just the first two weeks.")

h3("Adoption Metrics: What to Track and When")
body("Track adoption metrics weekly during the first eight weeks after launch.")

table_ncol(
    ["Metric", "What it tells you", "When to act"],
    [
        ("Active usage rate\n(% of eligible users who used the tool this week)",
         "Whether people are actually using the tool, not just trained on it",
         "Below 60% after week 2: investigate with a direct conversation, not a survey"),
        ("Repeat usage rate\n(% of week 1 users still using in week 4)",
         "Whether usage is becoming habit or a one-time trial",
         "Drop greater than 15% from week 1 to week 4: identify what is driving the drop-off"),
        ("Workaround rate\n(% of eligible tasks still done the old way)",
         "Whether the old process has been retired",
         "Any workaround above 10%: identify the specific cases driving it and address them explicitly"),
        ("User satisfaction\n(Would you recommend this tool to a colleague?)",
         "Leading indicator of sustained adoption and advocacy",
         "Negative trend two weeks running: run listening sessions to understand the specific friction"),
        ("Feedback volume\n(Issues flagged through the in-tool feedback channel)",
         "Whether users trust the feedback channel and whether you are closing the loop",
         "Sudden drop in feedback often means users stopped expecting a response. Close the loop visibly."),
    ]
)

case_study(
    "MCC Speaker Agent: catching the adoption plateau before it became a loss",
    "After the Speaker Agent launched, the MCC team tracked active usage weekly. Week 1 opened at "
    "65%, climbed to 85% by week 2, and reached 90% by week 3. Then it dropped to 85% in weeks 4 "
    "and 5. Without weekly tracking, that drop would not have been noticed for a month. With it, "
    "the change owner investigated within days and found that 15% of counselors were still using "
    "the old manual checklist for 'complex cases.'",
    "The team ran targeted office hours showing how to handle those specific cases with the AI, "
    "had champion counselors reach out to holdouts directly, and the executive sponsor sent a "
    "message confirming the old checklist was retired. Adoption climbed back to 92% within a week "
    "and held there. The active monitoring and fast intervention — not the launch — produced "
    "sustained adoption."
)

case_study(
    "MCC Lesseps: monitoring beyond usage and sustaining the change",
    "The MCC Lesseps team tracked more than logins. They watched the quality and specificity of "
    "user feedback, how clearly stakeholders understood the system, and whether users could "
    "articulate what worked, what did not, and what was still unknown. Overrides were framed as "
    "learning signals, not failures: each one was a data point about where the system needed "
    "improvement. The question shifted from 'Is the AI right?' to 'How do we make the system "
    "better?' That reframe prevented the common trap of reading an adoption plateau as user "
    "resistance when it is actually a signal about the product.",
    "Post-launch, the team sustained specific behaviors rather than declaring victory at deployment: "
    "AI output treated as decision support with human review as the default, decisions made on "
    "documented criteria, and leaders modeling the new workflow and holding the line when pressure "
    "tempted people back to the old way."
)

h3("Artifacts")
table_ncol(
    ["Artifact", "Description"],
    [
        ("Change Management Plan",
         "Full plan covering personas, change impact, champion network, communications timeline, "
         "training approach, feedback channels, and adoption metrics. Built through Engage and "
         "Discovery, activated in Execute."),
        ("UAT Feedback Tracker",
         "Tracks real-user testing feedback by issue, severity, acceptance criteria, ownership, "
         "and resolution status. The feedback loop that turns an imperfect MVP into something "
         "users trust."),
        ("Adoption Dashboard",
         "Weekly tracking of active usage rate, repeat usage, workaround rate, user satisfaction, "
         "and feedback volume. The primary tool for spotting and responding to adoption plateaus early."),
    ]
)

note("Adoption sustained",
     "The domain owns the solution, the old process is retired, adoption is holding above target, "
     "and the change owner has a clear intervention playbook for when it plateaus. Change management "
     "succeeded because it started in Engage, not at launch.")

# ── Save ──────────────────────────────────────────────────────────────────────
out = "adoption-change.docx"
doc.save(out)
print(f"Saved: {out}")
